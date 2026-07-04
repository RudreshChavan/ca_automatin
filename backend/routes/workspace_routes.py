"""
Workspace routes — admin-controlled document hub.
Create, edit, upload, manage, and share Word documents with clients.
"""

import os
import shutil
from flask import Blueprint, request, jsonify, g, send_from_directory

from middleware import require_auth, require_admin
from models import (
    create_workspace_document, get_workspace_documents, get_workspace_document_by_id,
    update_workspace_document, delete_workspace_document, increment_workspace_version,
    create_workspace_share, get_workspace_shares, get_shared_workspace_docs,
    revoke_workspace_share, check_workspace_access, get_workspace_stats,
    get_employees_and_clients, create_notification,
)

workspace_bp = Blueprint("workspace", __name__, url_prefix="/api/workspace")

# Workspace storage directory
WORKSPACE_DIR = os.path.join(os.path.dirname(__file__), "..", "local_storage", "workspace")
os.makedirs(WORKSPACE_DIR, exist_ok=True)


# ══════════════════════════════════════════════════════════════
# DOCUMENT CRUD
# ══════════════════════════════════════════════════════════════

@workspace_bp.route("/documents", methods=["POST"])
@require_auth
@require_admin
def create_document():
    """
    Create a new blank Word document.
    Body: { title, description?, status? }
    """
    data = request.get_json()
    if not data or not data.get("title"):
        return jsonify({"error": "Title is required"}), 400

    title = data["title"].strip()
    description = data.get("description", "").strip()
    status = data.get("status", "draft")

    if status not in ("draft", "published"):
        return jsonify({"error": "Status must be 'draft' or 'published'"}), 400

    # Generate a sanitized filename
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    if not safe_title:
        safe_title = "Untitled"
    file_name = f"{safe_title}.docx"

    try:
        # Create a blank .docx using aspose.words
        import docx
        doc = docx.Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph("")

        # Save to a temporary ID-based directory (we'll update after DB insert)
        temp_dir = os.path.join(WORKSPACE_DIR, "_temp")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, file_name)
        doc.save(temp_path)
        file_size = os.path.getsize(temp_path)

        # Create DB record
        doc_record = create_workspace_document(
            title=title,
            description=description,
            file_name=file_name,
            file_size=file_size,
            file_path="",  # Will update after getting the ID
            created_by=g.user_id,
            status=status,
        )

        # Move to permanent location
        doc_dir = os.path.join(WORKSPACE_DIR, str(doc_record["id"]))
        os.makedirs(doc_dir, exist_ok=True)
        final_path = os.path.join(doc_dir, file_name)
        shutil.move(temp_path, final_path)

        # Update the file_path in DB
        relative_path = f"workspace/{doc_record['id']}/{file_name}"
        update_workspace_document(doc_record["id"], updated_by=g.user_id)

        # Update file_path directly
        from database import get_connection
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE workspace_documents SET file_path = %s WHERE id = %s",
                       (relative_path, doc_record["id"]))
        conn.commit()
        cursor.close()
        conn.close()

        doc_record["file_path"] = relative_path

        return jsonify({"message": "Document created", "document": doc_record}), 201

    except Exception as e:
        return jsonify({"error": f"Failed to create document: {str(e)}"}), 500


@workspace_bp.route("/documents/upload", methods=["POST"])
@require_auth
@require_admin
def upload_document():
    """
    Upload an existing Word document to the workspace.
    Accepts multipart/form-data with:
        - file: the .docx file
        - title (optional): override title
        - description (optional)
        - status (optional): 'draft' or 'published'
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    uploaded = request.files["file"]
    if uploaded.filename == "":
        return jsonify({"error": "No file selected"}), 400

    filename = uploaded.filename
    name_lower = filename.lower()

    # Validate file type
    if not (name_lower.endswith('.docx') or name_lower.endswith('.doc')):
        return jsonify({"error": "Only Word documents (.doc, .docx) are supported"}), 400

    title = request.form.get("title", "").strip()
    if not title:
        title = os.path.splitext(filename)[0]
    description = request.form.get("description", "").strip()
    status = request.form.get("status", "draft")

    try:
        # Read file size
        uploaded.seek(0, os.SEEK_END)
        file_size = uploaded.tell()
        uploaded.seek(0)

        # Save to temp first
        temp_dir = os.path.join(WORKSPACE_DIR, "_temp")
        os.makedirs(temp_dir, exist_ok=True)
        temp_path = os.path.join(temp_dir, filename)
        uploaded.save(temp_path)

        # Convert .doc to .docx if needed
        if name_lower.endswith('.doc') and not name_lower.endswith('.docx'):
            os.remove(temp_path)
            return jsonify({"error": "Legacy .doc format is no longer supported for uploads. Please convert your file to .docx before uploading."}), 400

        # Create DB record
        doc_record = create_workspace_document(
            title=title,
            description=description,
            file_name=filename,
            file_size=file_size,
            file_path="",
            created_by=g.user_id,
            status=status,
        )

        # Move to permanent location
        doc_dir = os.path.join(WORKSPACE_DIR, str(doc_record["id"]))
        os.makedirs(doc_dir, exist_ok=True)
        final_path = os.path.join(doc_dir, filename)
        shutil.move(temp_path, final_path)

        # Update file_path
        relative_path = f"workspace/{doc_record['id']}/{filename}"
        from database import get_connection
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE workspace_documents SET file_path = %s WHERE id = %s",
                       (relative_path, doc_record["id"]))
        conn.commit()
        cursor.close()
        conn.close()

        doc_record["file_path"] = relative_path

        return jsonify({"message": "Document uploaded", "document": doc_record}), 201

    except Exception as e:
        return jsonify({"error": f"Failed to upload document: {str(e)}"}), 500


@workspace_bp.route("/documents", methods=["GET"])
@require_auth
@require_admin
def list_documents():
    """List all workspace documents (admin). Query: ?status=draft&search=xyz"""
    status_filter = request.args.get("status", "").strip() or None
    search = request.args.get("search", "").strip() or None
    docs = get_workspace_documents(status_filter=status_filter, search=search)
    return jsonify({"documents": docs}), 200


@workspace_bp.route("/documents/<int:doc_id>", methods=["GET"])
@require_auth
def get_document(doc_id):
    """Get a single document. Admins get full access; clients need a share."""
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    if getattr(g, 'role', 'user') != 'admin':
        access = check_workspace_access(g.user_id, doc_id)
        if not access["has_access"]:
            return jsonify({"error": "Access denied"}), 403

    return jsonify({"document": doc}), 200


@workspace_bp.route("/documents/<int:doc_id>", methods=["PUT"])
@require_auth
@require_admin
def update_document(doc_id):
    """Update document metadata. Body: { title?, description?, status? }"""
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "Request body is required"}), 400

    title = data.get("title")
    description = data.get("description")
    status = data.get("status")

    if status and status not in ("draft", "published"):
        return jsonify({"error": "Status must be 'draft' or 'published'"}), 400

    update_workspace_document(
        doc_id, title=title, description=description,
        status=status, updated_by=g.user_id,
    )

    updated = get_workspace_document_by_id(doc_id)
    return jsonify({"message": "Document updated", "document": updated}), 200


@workspace_bp.route("/documents/<int:doc_id>/content", methods=["PUT"])
@require_auth
@require_admin
def update_document_content(doc_id):
    """
    Save edited document content (HTML from CKEditor → .docx).
    Accepts JSON: { content: "<html>..." }
    Or FormData with a 'file' field containing a .docx blob.
    """
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Resolve full file path
    file_path = doc["file_path"]
    full_path = os.path.join(os.path.dirname(__file__), "..", "local_storage", file_path)
    full_path = os.path.abspath(full_path)

    try:
        if "file" in request.files:
            # Binary upload (docx blob)
            uploaded_file = request.files["file"]
            uploaded_file.save(full_path)
        else:
            # HTML content → .docx conversion
            data = request.get_json()
            if not data or "content" not in data:
                return jsonify({"error": "No content provided"}), 400

            html_content = data["content"]

            import docx
            from htmldocx import HtmlToDocx
            doc = docx.Document()
            new_parser = HtmlToDocx()
            new_parser.add_html_to_document(html_content, doc)
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            doc.save(full_path)

        # Update metadata
        new_size = os.path.getsize(full_path)
        increment_workspace_version(doc_id, g.user_id, new_size)

        return jsonify({"message": "Document content saved", "new_size": new_size}), 200

    except Exception as e:
        return jsonify({"error": f"Failed to save content: {str(e)}"}), 500


@workspace_bp.route("/documents/<int:doc_id>/download", methods=["GET"])
@require_auth
def download_document(doc_id):
    """Download/serve a workspace document. Admins or shared users."""
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Access control
    if getattr(g, 'role', 'user') != 'admin':
        access = check_workspace_access(g.user_id, doc_id)
        if not access["has_access"]:
            return jsonify({"error": "Access denied"}), 403
        if access["permission"] not in ("view", "download"):
            return jsonify({"error": "Access denied"}), 403

    file_path = doc["file_path"]
    full_path = os.path.join(os.path.dirname(__file__), "..", "local_storage", file_path)
    full_path = os.path.abspath(full_path)

    if not os.path.exists(full_path):
        return jsonify({"error": "File not found on disk"}), 404

    directory = os.path.dirname(full_path)
    filename = os.path.basename(full_path)
    return send_from_directory(directory, filename, as_attachment=True)


@workspace_bp.route("/documents/<int:doc_id>", methods=["DELETE"])
@require_auth
@require_admin
def remove_document(doc_id):
    """Delete a workspace document and its file."""
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Delete file from disk
    doc_dir = os.path.join(WORKSPACE_DIR, str(doc_id))
    if os.path.isdir(doc_dir):
        shutil.rmtree(doc_dir, ignore_errors=True)

    # Delete from DB (cascades shares)
    delete_workspace_document(doc_id)

    return jsonify({"message": "Document deleted"}), 200


# ══════════════════════════════════════════════════════════════
# SHARING
# ══════════════════════════════════════════════════════════════

@workspace_bp.route("/documents/<int:doc_id>/share", methods=["POST"])
@require_auth
@require_admin
def share_document(doc_id):
    """
    Share a workspace document with a client/employee.
    Body: { shared_with_user_id, permission? }
    """
    doc = get_workspace_document_by_id(doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    data = request.get_json()
    if not data or not data.get("shared_with_user_id"):
        return jsonify({"error": "shared_with_user_id is required"}), 400

    shared_with = data["shared_with_user_id"]
    permission = data.get("permission", "view")

    if permission not in ("view", "download"):
        return jsonify({"error": "Permission must be 'view' or 'download'"}), 400

    try:
        share = create_workspace_share(
            document_id=doc_id,
            shared_by=g.user_id,
            shared_with=shared_with,
            permission=permission,
        )

        # Notify the recipient
        try:
            perm_label = {"view": "View Only", "download": "Download"}
            create_notification(
                user_id=shared_with,
                notif_type="workspace_shared",
                title="Document Shared With You",
                message=f"'{doc['title']}' has been shared with you ({perm_label.get(permission, permission)} access).",
                related_id=doc_id,
            )
        except Exception as e:
            print(f"[Notify] Failed to notify workspace share: {e}")

        return jsonify({"message": "Document shared", "share": share}), 201

    except Exception as e:
        error_str = str(e)
        if "uq_wsshare" in error_str.lower() or "unique" in error_str.lower():
            return jsonify({"error": "This document is already shared with that user."}), 409
        return jsonify({"error": f"Failed to share: {error_str}"}), 500


@workspace_bp.route("/documents/<int:doc_id>/shares", methods=["GET"])
@require_auth
@require_admin
def list_document_shares(doc_id):
    """List all shares for a workspace document."""
    shares = get_workspace_shares(doc_id)
    return jsonify({"shares": shares}), 200


@workspace_bp.route("/shares/<int:share_id>", methods=["DELETE"])
@require_auth
@require_admin
def remove_share(share_id):
    """Revoke a workspace share."""
    success = revoke_workspace_share(share_id)
    if not success:
        return jsonify({"error": "Share not found"}), 404
    return jsonify({"message": "Share revoked"}), 200


@workspace_bp.route("/shared", methods=["GET"])
@require_auth
def my_shared_documents():
    """Get workspace documents shared with the current user (client view)."""
    docs = get_shared_workspace_docs(g.user_id)
    return jsonify({"documents": docs}), 200


@workspace_bp.route("/users", methods=["GET"])
@require_auth
@require_admin
def shareable_users():
    """Get list of employees and clients for the share dropdown."""
    users = get_employees_and_clients()
    return jsonify({"users": users}), 200


# ══════════════════════════════════════════════════════════════
# STATISTICS
# ══════════════════════════════════════════════════════════════

@workspace_bp.route("/stats", methods=["GET"])
@require_auth
@require_admin
def workspace_statistics():
    """Get workspace statistics."""
    stats = get_workspace_stats()
    return jsonify({"stats": stats}), 200
